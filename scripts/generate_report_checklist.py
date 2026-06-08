# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

title = doc.add_heading('Чек-лист отчёта по проекту', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Проект: Assist Eye')
doc.add_paragraph('Мобильное приложение для помощи слабовидящим и незрячим (голосовое управление, OCR, детекция объектов, распознавание купюр, навигация; офлайн-режим).')
doc.add_paragraph()

def add_field(title_text, body_paragraphs):
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    for text in body_paragraphs:
        doc.add_paragraph(text)

add_field('ФИО (проект):', [
    '[Укажите ваше ФИО]',
    'Проект: Assist Eye',
])

add_field('Что удалось сделать / завершить (артефакты, задачи, ссылки):', [
    '',
    '1. Системное тестирование',
    '• Настроена среда сборки и запуска: npm install, Expo dev build (expo run:android), конфигурация .env, eas.json.',
    '• Прогон автоматических проверок: npm run validate (typecheck, lint, jest) — unit-тесты для ядра (result, errorCodes, featureFlags, appConstants, StartupPermissions).',
    '• Системное тестирование сборки Android: устранены блокеры Gradle/SDK (NDK, CMake, Build Tools, minSdk 26 для Yandex MapKit, TLS/репозитории, patch-package для react-native-localize, react-native-permissions, react-native-worklets-core, @bam.tech/react-native-image-resizer).',
    '• Выявлены ограничения: приложение не работает в Expo Go (нужен dev build из-за нативных модулей: Vosk, Vision Camera, TFLite, Yandex MapKit); iOS — только на macOS или через EAS Build.',
    '• Сборка Android дошла до стадии компиляции нативных модулей (~570 Gradle-задач); финальная сборка APK требует завершения после исправления package name / autolinking.',
    '',
    '2. Доработка основной функциональности',
    '• Восстановлены отсутствующие слои data/*: хранилище (databaseHelper, репозитории настроек и истории), TTS (reactNativeTtsService на expo-speech), камера (cameraService), распознавание речи (заглушка Vosk), навигация (Yandex), фабрика зависимостей commandProcessor.',
    '• Сохранена архитектура: commandStore (голосовые команды: чтение, описание сцены, купюра, навигация, настройки), CameraHost + frame processor, i18n (RU/EN), экраны Main / Settings / History.',
    '• Интеграции в проекте: ML Kit OCR, TFLite (детекция/купюры), Vosk, Vision Camera, Yandex MapKit, expo-sqlite, feature flags.',
    '',
    'Ссылка на результат: репозиторий assist-eye [указать URL GitHub, если выложен].',
])

add_field('Субъективная завершённость проекта (готовность к защите):', [
    'Оценка: 65–75%',
    '',
    'Обоснование: архитектура и UI в целом реализованы; часть data-реализаций — заглушки; Android-сборка близка к успеху, но стабильный «зелёный» билд на устройстве не зафиксирован; полноценное E2E/системное тестирование на реальном железе и демо всех сценариев — в процессе.',
])

add_field('Ожидания в начале семестра vs факт:', [
    '• Ожидание: полностью офлайн приложение с голосом, OCR, CV, купюрами → Факт: заложено в код и модели; часть сервисов пока stub → Частично совпало.',
    '• Ожидание: быстрый запуск через Expo Go → Факт: нужен native dev build → Не совпало.',
    '• Ожидание: стабильная сборка Android/iOS → Факт: длительная отладка Android; iOS не проверялся локально (Windows) → Частично совпало.',
    '• Ожидание: системное тестирование всех модулей → Факт: unit-тесты ядра + отладка сборки; сценарное тестирование на устройстве не завершено → Частично совпало.',
    '',
    'Причина расхождений: стек с множеством нативных библиотек (камера, TFLite, Vosk, карты) усложнил сборку; значительная часть времени ушла на инфраструктуру, а не только на бизнес-логику.',
])

add_field('Презентация, отчёт, стратегия защиты:', [
    '• Презентация: [в процессе / не начата — указать актуальный статус].',
    '• Отчёт: в процессе (данный чек-лист — часть отчётных материалов).',
    '• Стратегия защиты: демонстрация на dev build Android; запасной план — видео/скринкаст; акцент на архитектуру, офлайн-режим, доступность (a11y).',
])

doc.add_paragraph()
doc.add_heading('Краткая версия для формы', level=2)
doc.add_paragraph(
    'Assist Eye — RN/Expo-приложение для слабовидящих (голос, OCR, детекция, купюры, карты, офлайн). '
    'Выполнены: настройка dev-сборки Android, unit-тесты ядра, устранение блокеров Gradle/SDK, восстановление data-слоя. '
    'Завершённость ~70%. Ожидали простой Expo Go — фактически потребовался native build. '
    'Презентация/отчёт: в работе.'
)

out = r'c:\assist-eye\Отчет_чеклист_Assist_Eye.docx'
doc.save(out)
print(out)
